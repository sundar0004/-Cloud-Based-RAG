from pathlib import Path

from docx import Document
from docx.shared import Pt

SRC = Path("docs/FINAL_SUBMISSION_DOCUMENT.md")
OUT = Path("docs/FINAL_SUBMISSION_DOCUMENT.docx")
ARCH_IMG = Path("docs/figures/architecture_diagram.png")
GRAPH_IMG = Path("docs/figures/results_graph.png")
FRONTEND_IMG = Path("docs/screenshots/frontend_ui.png")
BACKEND_IMG = Path("docs/screenshots/backend_api_docs.png")


def add_code_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def main():
    if not SRC.exists():
        raise FileNotFoundError(f"Missing source markdown: {SRC}")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    lines = SRC.read_text(encoding="utf-8").splitlines()
    in_code = False
    first_h1 = True

    for raw in lines:
        line = raw.rstrip()

        if line.strip().startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            if line.strip():
                add_code_paragraph(doc, line)
            else:
                doc.add_paragraph("")
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            if not first_h1:
                doc.add_page_break()
            first_h1 = False
            doc.add_heading(line[2:].strip(), level=1)
            continue

        if line.startswith("## "):
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=2)
            if "High-Level Architecture Diagram" in heading_text and ARCH_IMG.exists():
                doc.add_picture(str(ARCH_IMG), width=Pt(500))
            if "Results and Analysis" in heading_text and GRAPH_IMG.exists():
                doc.add_picture(str(GRAPH_IMG), width=Pt(500))
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        if line.startswith("- ") or line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. "):
            p = doc.add_paragraph(line)
            continue

        if line.startswith("|"):
            add_code_paragraph(doc, line)
            continue

        doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("Performance Graph Evidence", level=1)
    if GRAPH_IMG.exists():
        doc.add_paragraph("Figure: Baseline vs RAG performance graph")
        doc.add_picture(str(GRAPH_IMG), width=Pt(500))

    doc.add_page_break()
    doc.add_heading("Application UI Evidence", level=1)
    if FRONTEND_IMG.exists():
        doc.add_paragraph("Figure: Frontend UI")
        doc.add_picture(str(FRONTEND_IMG), width=Pt(500))
    if BACKEND_IMG.exists():
        doc.add_paragraph("Figure: Backend Swagger UI")
        doc.add_picture(str(BACKEND_IMG), width=Pt(500))

    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    main()
